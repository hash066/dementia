import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

doc = docx.Document()

# Title
title = doc.add_heading('Dementor: Autonomous Edge-AI for Proactive Dementia Care', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_heading('A Zero-Trust, Tri-Node Architecture Fusing Gemma 4 RAG and On-Device LiteRT for Secure Elder Monitoring', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Track
p = doc.add_paragraph()
p.add_run('Selected Track: ').bold = True
p.add_run('Edge AI / AI for Social Good')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('---')

# Section 1
add_heading(doc, '1. The Vision & The Problem', 2)
p = doc.add_paragraph('Dementia care requires vigilant monitoring, but traditional solutions force families into an uncomfortable compromise: sacrifice privacy by uploading constant audio/video to the cloud, or risk safety by relying on manual check-ins. Furthermore, cloud-dependent systems are brittle; a WiFi outage during a fall emergency can be fatal.')

p = doc.add_paragraph()
p.add_run('Dementor').bold = True
p.add_run(' is our answer. Designed specifically for the Gemma 4 Hackathon, it is a completely autonomous, offline-first elder care monitoring system. By pushing intelligence to the edge, we achieved a system that listens, sees, and analyzes in real-time, without a single byte of sensitive medical data ever leaving the local network. Our primary engineering goal was to prove that complex AI workflows can run entirely on consumer edge hardware, unlocking the true potential of Gemma 4 for high-stakes healthcare applications.')

# Section 2
add_heading(doc, '2. The Tri-Node Architecture', 2)
doc.add_paragraph('To achieve zero-latency monitoring and complete privacy, we decoupled our architecture into three distinct layers, integrated via strict Protocol Buffer contracts.')

add_heading(doc, 'Node 1: The Sensory Edge (ESP32)', 3)
doc.add_paragraph('We built custom C-firmware for ESP32 microcontrollers. It handles high-frequency sensory ingestion with strict real-time constraints:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Audio Pipeline: ').bold = True
p.add_run('I2S capture from PDM microphones, processed into 20ms frames, Opus-encoded, and streamed over RTP/UDP to avoid TCP overhead.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Inertial Pipeline: ').bold = True
p.add_run('I2C MPU6050 sampling at 100Hz with a bespoke embedded finite state machine (FSM) for immediate fall detection (detecting 0.3g free-falls followed by >2.5g impact spikes).')

add_heading(doc, 'Node 2: The Intelligence Runtime (Raspberry Pi)', 3)
doc.add_paragraph('The Pi acts as the central local router and rapid-response unit, running 4 concurrent asynchronous pipelines:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Vision & Audio Processing: ').bold = True
p.add_run('Real-time OpenCV background subtraction feeds into a MobileNet SSD for patient and object tracking. Audio is buffered by a Silero VAD (ONNX) and transcribed locally using whisper.cpp.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Autonomous Emergency FSM: ').bold = True
p.add_run('If the Pi detects an anomaly (e.g., a fall + distress keywords) and the Phone Hub is disconnected, it acts autonomously—escalating priority and triggering local alarms via Piper TTS.')

add_heading(doc, 'Node 3: The Caregiver Hub & Jetpack Compose UI', 3)
doc.add_paragraph('The caregiver’s Android phone acts as the system\'s brain and primary interface. We built a native Android app using Kotlin and Jetpack Compose, communicating with a decoupled Python FastAPI backend running directly on the local device network.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('LiteRT Intent Classification: ').bold = True
p.add_run('For instantaneous app UI interactions, we embedded TensorFlow Lite (LiteRT) directly into the Android app to classify caregiver intents locally, avoiding unnecessary backend overhead.')

# Section 3
add_heading(doc, '3. Pushing the Limits: How We Leveraged Gemma 4', 2)
doc.add_paragraph('The core differentiator of Dementor is providing caregivers with a natural language interface to query the patient\'s status over time. Sending this data to a cloud LLM was a non-starter due to HIPAA and privacy concerns.')
doc.add_paragraph('Our solution was deploying the ').add_run('Gemma 4 model locally').bold = True
doc.paragraphs[-1].add_run(' on the Phone Hub (utilizing 4-bit GGUF quantization via llama.cpp). Gemma 4 is the beating heart of our system, utilized in two critical, concurrent pipelines:')

add_heading(doc, '1. Domain-Specific Fine-Tuning & Continuous Entity Extraction', 3)
p = doc.add_paragraph('Gemma 4 doesn\'t just answer questions; it acts as an asynchronous medical agent. Out of the box, LLMs can be overly conversational. We fine-tuned our Gemma 4 model using synthetic blends based on ')
p.add_run('MIMIC-IV, DementiaBank, and HealthCareMagic').bold = True
p.add_run(' datasets. When the Pi sends a SPEECH event, Gemma 4 processes the transcript in the background, extracting JSON-structured entities (e.g., {"category": "medication", "label": "aspirin", "distress": "mild"}) and generating a rolling summary. This tailored Gemma 4 to excel at elder-care specific sentiment analysis and recognizing subtle signs of cognitive decline.')

add_heading(doc, '2. Retrieval-Augmented Generation (RAG) Memory', 3)
doc.add_paragraph('We built an encrypted SQLCipher database wrapped in an SQLite FTS5 (Full-Text Search) index. When a caregiver asks the app, "Did my father mention his chest hurting today?"')
doc.add_paragraph('1. The backend performs a BM25 semantic search across the day\'s local ASR transcripts.', style='List Number')
doc.add_paragraph('2. The relevant snippets are injected into Gemma 4\'s context window.', style='List Number')
doc.add_paragraph('3. Gemma 4 streams a synthesized, clinically accurate, and deeply empathetic response back to the Jetpack Compose UI via Server-Sent Events (SSE).', style='List Number')
doc.add_paragraph('By leveraging Gemma 4\'s incredible context understanding and fast local generation, we created a smart hub that feels as responsive as a cloud API but retains 100% data sovereignty.')

# Section 4
add_heading(doc, '4. Engineering Challenges Overcome', 2)
p = doc.add_paragraph()
p.add_run('Challenge 1: Real-Time Audio Inference on Edge Hardware\n').bold = True
p.add_run('Running continuous ASR on a Raspberry Pi originally choked the CPU, causing massive queues and delayed alerts.\n')
p.add_run('Solution: ').bold = True
p.add_run('We introduced a Silero VAD (Voice Activity Detection) gating mechanism. Instead of streaming endless room noise to Whisper, the VAD only captures active speech chunks, flushing them to whisper-base.en (147MB) only when silence is detected. This reduced CPU load by 85%.')

p = doc.add_paragraph()
p.add_run('Challenge 2: The Asynchronous Data Race\n').bold = True
p.add_run('With ESP32 fall events, Pi vision detections, and ASR transcripts all hitting the Phone Hub simultaneously, we faced race conditions resulting in duplicated alerts and corrupted emergency states.\n')
p.add_run('Solution: ').bold = True
p.add_run('We implemented an immutable, event-sourcing architecture. Every event from the Pi is packaged in an EventEnvelope Protobuf with a strict timestamp and UUIDv4. The FastAPI intake server passes them through an in-memory Bloom Filter to drop duplicates instantly, followed by a priority queue scheduler that enforces single-writer atomic transactions to the SQLCipher database.')

p = doc.add_paragraph()
p.add_run('Challenge 3: Rapid Multi-Developer Collaboration without Breakage\n').bold = True
p.add_run('Building hardware, AI pipelines, and an Android app simultaneously typically results in integration hell.\n')
p.add_run('Solution: ').bold = True
p.add_run('We adopted a strict "Zero Merge Conflict" policy. Every directory had a specific CODEOWNER. The edges between our domains (e.g., App to Backend, Pi to Phone) were frozen as OpenAPI specs and Protobuf definitions in a contracts/ directory in Week 1. This allowed the Android UI team to mock the Gemma 4 RAG endpoints perfectly while the backend team optimized the local inference.')

# Section 5
add_heading(doc, '5. Why Our Technical Choices Were Right', 2)
p = doc.add_paragraph(style='List Number')
p.add_run('Privacy as a Physical Guarantee: ').bold = True
p.add_run('By using Gemma 4 locally alongside whisper.cpp, we don\'t just promise privacy—we physically guarantee it. The system functions flawlessly without an outbound internet connection.')
p = doc.add_paragraph(style='List Number')
p.add_run('Resilient Local Networking: ').bold = True
p.add_run('Cloud solutions fail when ISPs fail. By operating entirely on mDNS and local network routing, Dementor guarantees that critical alerts (like a detected fall) reach the caregiver\'s phone in milliseconds.')
p = doc.add_paragraph(style='List Number')
p.add_run('Jetpack Compose + LiteRT Integration: ').bold = True
p.add_run('Building the app natively with Compose allowed us to create a premium, reactive user experience. Integrating LiteRT directly into the frontend offloaded simple intent routing from the backend, keeping the Python API free to handle the heavy-lifting Gemma 4 inference.')

doc.add_paragraph('We didn\'t just build an AI wrapper. We built a resilient, production-ready IoT edge network, secured it with modern cryptography, and empowered it with state-of-the-art open models. Dementor proves that with Gemma 4, the future of medical AI doesn\'t live in the cloud—it lives safely at home.')

doc.save('e:\\dementia\\Dementor_Gemma4_Hackathon_Writeup.docx')
print("Saved to e:\\dementia\\Dementor_Gemma4_Hackathon_Writeup.docx")
